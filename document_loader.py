"""
Document loader module for handling multiple file formats.
Supports: PDF, Word (.docx), TXT, Markdown (.md)
"""
import logging
import os
from pathlib import Path
from typing import Optional, Dict, Any
from enum import Enum

logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """Supported document formats."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    MD = "md"
    MARKDOWN = "markdown"


class DocumentLoader:
    """
    Universal document loader that handles multiple file formats.
    """
    
    def __init__(self):
        """Initialize the document loader."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are installed."""
        try:
            import pypdf
            self.has_pdf = True
        except ImportError:
            self.has_pdf = False
            logger.warning("pypdf not installed. PDF support disabled.")
        
        try:
            import docx
            self.has_docx = True
        except ImportError:
            self.has_docx = False
            logger.warning("python-docx not installed. Word support disabled.")
    
    def detect_format(self, file_path: str) -> Optional[DocumentFormat]:
        """
        Detect document format from file extension.
        
        Args:
            file_path: Path to the document
            
        Returns:
            DocumentFormat enum or None if unsupported
        """
        extension = Path(file_path).suffix.lower().lstrip('.')
        
        format_map = {
            'pdf': DocumentFormat.PDF,
            'docx': DocumentFormat.DOCX,
            'doc': DocumentFormat.DOC,
            'txt': DocumentFormat.TXT,
            'md': DocumentFormat.MD,
            'markdown': DocumentFormat.MARKDOWN,
        }
        
        return format_map.get(extension)
    
    def load_pdf(self, file_path: str) -> str:
        """
        Load content from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        if not self.has_pdf:
            raise ImportError(
                "pypdf is required for PDF support. "
                "Install with: pip install pypdf"
            )
        
        try:
            import pypdf
            
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                logger.info(f"Reading PDF with {num_pages} pages...")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                    
                    if page_num % 10 == 0:
                        logger.info(f"Processed {page_num}/{num_pages} pages")
            
            content = "\n\n".join(text_content)
            logger.info(f"Extracted {len(content)} characters from PDF")
            
            return content
            
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
    
    def load_docx(self, file_path: str) -> str:
        """
        Load content from Word (.docx) file.
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Extracted text content
        """
        if not self.has_docx:
            raise ImportError(
                "python-docx is required for Word support. "
                "Install with: pip install python-docx"
            )
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
            
            # Combine all content
            content = "\n\n".join(paragraphs)
            if table_text:
                content += "\n\n" + "\n".join(table_text)
            
            logger.info(f"Extracted {len(content)} characters from Word document")
            logger.info(f"Found {len(paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            return content
            
        except Exception as e:
            logger.error(f"Error loading Word document {file_path}: {e}")
            raise
    
    def load_txt(self, file_path: str) -> str:
        """
        Load content from text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            logger.info(f"Loaded {len(content)} characters from text file")
            return content
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                logger.warning(f"Loaded {file_path} with latin-1 encoding")
                return content
            except Exception as e:
                logger.error(f"Error loading text file {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            raise
    
    def load_markdown(self, file_path: str) -> str:
        """
        Load content from Markdown file.
        
        Args:
            file_path: Path to Markdown file
            
        Returns:
            File content (raw markdown)
        """
        # Markdown is just text, but we could optionally strip markdown syntax
        return self.load_txt(file_path)
    
    def load_document(self, file_path: str) -> str:
        """
        Load document content automatically detecting format.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Detect format
        doc_format = self.detect_format(file_path)
        
        if not doc_format:
            raise ValueError(
                f"Unsupported file format: {Path(file_path).suffix}. "
                f"Supported formats: PDF, DOCX, TXT, MD"
            )
        
        logger.info(f"Loading {doc_format.value.upper()} file: {file_path}")
        
        # Load based on format
        if doc_format == DocumentFormat.PDF:
            return self.load_pdf(file_path)
        
        elif doc_format in (DocumentFormat.DOCX, DocumentFormat.DOC):
            return self.load_docx(file_path)
        
        elif doc_format == DocumentFormat.TXT:
            return self.load_txt(file_path)
        
        elif doc_format in (DocumentFormat.MD, DocumentFormat.MARKDOWN):
            return self.load_markdown(file_path)
        
        else:
            raise ValueError(f"Unsupported format: {doc_format}")
    
    def get_supported_formats(self) -> Dict[str, bool]:
        """
        Get list of supported formats and their availability.
        
        Returns:
            Dictionary of format: available
        """
        return {
            "PDF (.pdf)": self.has_pdf,
            "Word (.docx)": self.has_docx,
            "Text (.txt)": True,
            "Markdown (.md)": True,
        }


# Global instance
document_loader = DocumentLoader()


def load_document_content(file_path: str) -> str:
    """
    Convenience function to load document content.
    
    Args:
        file_path: Path to document
        
    Returns:
        Extracted text content
    """
    return document_loader.load_document(file_path)


if __name__ == "__main__":
    # Test the loader
    logging.basicConfig(level=logging.INFO)
    
    loader = DocumentLoader()
    
    print("\n📄 Document Loader - Supported Formats:")
    print("-" * 50)
    for format_name, available in loader.get_supported_formats().items():
        status = "✅ Available" if available else "❌ Not installed"
        print(f"{format_name}: {status}")
    
    print("\n💡 To enable all formats, install:")
    print("   pip install pypdf python-docx")
